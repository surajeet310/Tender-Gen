import os
import pickle
import pythoncom
import win32com.client
from docx import Document
from docxcompose.composer import Composer
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from datetime import datetime
from os.path import exists


class TenderGen:
    def __init__(self, app_obj, pdf):
        self.app_obj = app_obj
        self.choice = pdf
        self.status = None
        self.path = None
        self.date = None
        self.merged_doc_file_name = (
            "All Tender Declaration"
            + "_"
            + self.app_obj.tender_name_of_work.GetValue()
            + ".docx"
        )
        self.work_duration = list()
        self.bid_validity = list()
        self.list_of_prefs = list()
        self.get_local_date()
        self.init_doc()
        self.get_location()
        self.get_order_prefs()
        self.get_work_duration()
        self.get_bid_validity()

    def get_local_date(self):
        now = datetime.now()
        self.date = now.strftime("%d/%m/%Y")

    def get_location(self):
        if exists("./location_pref.pickle"):
            loc_pref = open("./location_pref.pickle", "rb")
            loc_pref_dict = pickle.load(loc_pref)
            self.path = loc_pref_dict["location"]
            loc_pref.close()
        else:
            self.path = os.getcwd() + "/"

    def get_work_duration(self):
        if int(self.app_obj.tender_duration_years.GetValue()) != 0:
            self.work_duration.append(
                self.app_obj.tender_duration_years.GetValue() + " Years"
            )
        if int(self.app_obj.tender_duration_months.GetValue()) != 0:
            self.work_duration.append(
                self.app_obj.tender_duration_months.GetValue() + " Months"
            )
        if int(self.app_obj.tender_duration_days.GetValue()) != 0:
            self.work_duration.append(
                self.app_obj.tender_duration_days.GetValue() + " Days"
            )

    def get_bid_validity(self):
        if int(self.app_obj.bid_validity_years.GetValue()) != 0:
            self.bid_validity.append(
                self.app_obj.bid_validity_years.GetValue() + " Years"
            )
        if int(self.app_obj.bid_validity_months.GetValue()) != 0:
            self.bid_validity.append(
                self.app_obj.bid_validity_months.GetValue() + " Months"
            )
        if int(self.app_obj.bid_validity_days.GetValue()) != 0:
            self.bid_validity.append(
                self.app_obj.bid_validity_days.GetValue() + " Days"
            )

    def get_order_prefs(self):
        if exists("./order_pref.pickle"):
            order_pref_file = open("./order_pref.pickle", "rb")
            order_pref = pickle.load(order_pref_file)
            for i in range(1, 16):
                if order_pref[str(i)] == "Qualification Information":
                    self.list_of_prefs.append("qualification_info_gen.docx")
                if order_pref[str(i)] == "Tender Details":
                    self.list_of_prefs.append("tender_details_gen.docx")
                if order_pref[str(i)] == "Bidder Details":
                    self.list_of_prefs.append("bidder_details_gen.docx")
                if order_pref[str(i)] == "Methodology of work":
                    self.list_of_prefs.append("methodology_work_gen.docx")
                if order_pref[str(i)] == "Undertaking(No objection)":
                    self.list_of_prefs.append("undertaking_no_objection_gen.docx")
                if order_pref[str(i)] == "Undertaking(On going work)":
                    self.list_of_prefs.append("undertaking_on_going_work_gen.docx")
                if order_pref[str(i)] == "Undertaking(Bid validity)":
                    self.list_of_prefs.append("undertaking_bid_validity_gen.docx")
                if order_pref[str(i)] == "Undertaking(Min. cash invest)":
                    self.list_of_prefs.append("undertaking_min_cash_gen.docx")
                if order_pref[str(i)] == "Declaration(No near relatives)":
                    self.list_of_prefs.append("declaration_no_relatives_gen.docx")
                if order_pref[str(i)] == "Authorization seek reference":
                    self.list_of_prefs.append("auth_seek_ref_gen.docx")
                if (
                    order_pref[str(i)]
                    == "Acceptance/Non acceptance of dispute review expert"
                ):
                    self.list_of_prefs.append(
                        "acceptance_review_dispute_expert_gen.docx"
                    )
                if (
                    order_pref[str(i)]
                    == "Information on Litigation History in which the bidder is involved"
                ):
                    self.list_of_prefs.append("docs/info_litigation_history.docx")
                if (
                    order_pref[str(i)]
                    == "Proposed sub-contractors and firms involved for construction"
                ):
                    self.list_of_prefs.append(
                        "docs/proposed_sub_contractors_for_const.docx"
                    )
                if order_pref[str(i)] == "Works for which bids are already submitted":
                    self.list_of_prefs.append("docs/works_bids_submitted.docx")
                if order_pref[str(i)] == "Milestone":
                    self.list_of_prefs.append("milestone_gen.docx")
            order_pref_file.close()
        else:
            self.list_of_prefs = [
                "qualification_info_gen.docx",
                "tender_details_gen.docx",
                "bidder_details_gen.docx",
                "methodology_work_gen.docx",
                "undertaking_no_objection_gen.docx",
                "undertaking_on_going_work_gen.docx",
                "undertaking_bid_validity_gen.docx",
                "undertaking_min_cash_gen.docx",
                "declaration_no_relatives_gen.docx",
                "auth_seek_ref_gen.docx",
                "acceptance_review_dispute_expert_gen.docx",
                "docs/info_litigation_history.docx",
                "docs/proposed_sub_contractors_for_const.docx",
                "docs/works_bids_submitted.docx",
                "milestone_gen.docx",
            ]

    def init_doc(self):
        self.qualification_info = Document("docs/qualification_info.docx")
        self.tender_info = Document("docs/tender_details.docx")
        self.bidder_info = Document("docs/bidder_details.docx")
        self.methodology_info = Document("docs/methodology_work.docx")
        self.undertaking_no_objection_info = Document(
            "docs/undertaking_no_objection.docx"
        )
        self.undertaking_on_going_work_info = Document(
            "docs/undertaking_on_going_work.docx"
        )
        self.undertaking_bid_validity_info = Document(
            "docs/undertaking_bid_validity.docx"
        )
        self.undertaking_min_cash_info = Document("docs/undertaking_min_cash.docx")
        self.auth_seek_ref_info = Document("docs/auth_seek_ref.docx")
        self.declaration_info = Document("docs/declaration_no_relatives.docx")
        self.dispute_review_info = Document(
            "docs/acceptance_review_dispute_expert.docx"
        )
        self.proposed_sub_contractors_info = Document(
            "docs/proposed_sub_contractors_for_const.docx"
        )
        self.works_bids_submitted_info = Document("docs/works_bids_submitted.docx")
        self.litigation_history_info = Document("docs/info_litigation_history.docx")
        self.milestone_info = Document("docs/milestone.docx")

    def generate(self):
        self.qualification_info_doc_process()
        self.tender_details_doc_process()
        self.bidder_details_doc_process()
        self.methodology_doc_process()
        self.undertaking_no_objection_doc_process()
        self.undertaking_on_going_work_doc_process()
        self.undertaking_bid_validity_doc_process()
        self.undertaking_cash_invest_doc_process()
        self.declaration_doc_process()
        self.auth_seek_ref_doc_process()
        self.dispute_review_doc_process()
        self.milestone_doc_process()
        self.combine_doc()
        if self.choice:
            self.pdf_generate()

    def qualification_info_doc_process(self):
        count = 5
        for paragraph in self.qualification_info.paragraphs:
            if count == 0:
                break
            if "<bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_name>",
                            self.app_obj.bidder_name.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<address_of_bidder>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<address_of_bidder>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<address_of_bidder>",
                            self.app_obj.bidder_address.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<place_of_reg>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<place_of_reg>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<place_of_reg>",
                            self.app_obj.place_of_reg.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<place_of_buisness>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<place_of_buisness>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<place_of_buisness>",
                            self.app_obj.place_of_buisness.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidding_type>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidding_type>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidding_type>",
                            self.app_obj.bidding_type.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
        self.qualification_info.save("qualification_info_gen.docx")

    def tender_details_doc_process(self):
        index = None
        tags = [
            "<tender_ref_num>",
            "<name_of_work>",
            "<pckg_num>",
            "<employer_name>",
            "<estm_cost>",
            "<earnest_money>",
            "<cost_of_paper>",
            "<completion_duration>",
        ]
        tags_values = [
            self.app_obj.tender_ref_num.GetValue(),
            self.app_obj.tender_name_of_work.GetValue(),
            self.app_obj.tender_pckg_num.GetValue(),
            self.app_obj.tender_employer_name.GetValue(),
            self.app_obj.tender_estm_cost.GetValue(),
            self.app_obj.tender_earnest_money.GetValue(),
            self.app_obj.tender_paper_cost.GetValue(),
            " ".join(self.work_duration),
        ]
        for row in self.tender_info.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for tag in range(len(tags)):
                        if tags[tag] in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if tags[tag] in inline[i].text:
                                    new_text = inline[i].text.replace(
                                        tags[tag],
                                        tags_values[tag],
                                    )
                                    inline[i].text = new_text
                                    index = i
                        else:
                            index = None
                    if (len(tags) != 0) & (index != None):
                        del tags[index]
                        del tags_values[index]

        for p in self.tender_info.paragraphs:
            if "<date>" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace("<date>", self.date)
                        inline[i].text = new_text
        self.tender_info.save("tender_details_gen.docx")

    def bidder_details_doc_process(self):
        index = None
        tags = [
            "<name_of_contractor>",
            "<bidding_type>",
            "<reg_num>",
            "<address_of_bidder>",
            "<bidder_ph_num>",
            "<bidder_email_id>",
        ]
        tags_values = [
            self.app_obj.contractor_name.GetValue(),
            self.app_obj.bidding_type.GetValue(),
            self.app_obj.reg_num.GetValue(),
            self.app_obj.bidder_address.GetValue(),
            self.app_obj.bidder_ph_num.GetValue(),
            self.app_obj.bidder_email_id.GetValue(),
        ]
        for row in self.bidder_info.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for tag in range(len(tags)):
                        if tags[tag] in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if tags[tag] in inline[i].text:
                                    new_text = inline[i].text.replace(
                                        tags[tag],
                                        tags_values[tag],
                                    )
                                    inline[i].text = new_text
                                    index = i
                        else:
                            index = None
                    if (len(tags) != 0) & (index != None):
                        del tags[index]
                        del tags_values[index]

        for p in self.bidder_info.paragraphs:
            if "<date>" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace("<date>", self.date)
                        inline[i].text = new_text
        self.bidder_info.save("bidder_details_gen.docx")

    def methodology_doc_process(self):
        count = 2
        for paragraph in self.methodology_info.paragraphs:
            if count == 0:
                break
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<completion_duration>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<completion_duration>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<completion_duration>", " ".join(self.work_duration)
                        )
                        inline[i].text = new_text
                        count -= 1
        self.methodology_info.save("methodology_work_gen.docx")

    def undertaking_no_objection_doc_process(self):
        for p in self.undertaking_no_objection_info.paragraphs:
            if "<date>" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace("<date>", self.date)
                        inline[i].text = new_text
                break
        self.undertaking_no_objection_info.save("undertaking_no_objection_gen.docx")

    def undertaking_on_going_work_doc_process(self):
        count = 6
        for paragraph in self.undertaking_on_going_work_info.paragraphs:
            if count == 0:
                break
            if "<name_of_contractor>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_contractor>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_contractor>",
                            '"' + self.app_obj.contractor_name.GetValue() + '"',
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_address>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_address>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_address>",
                            self.app_obj.bidder_address.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<employer_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<employer_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<employer_name>",
                            self.app_obj.tender_employer_name.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<temp_bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<temp_bidder_name>" in inline[i].text:
                        if self.app_obj.bidding_type.GetValue() == "Partnership":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "Authorised Signatory and Partner of "
                                + self.app_obj.bidder_name.GetValue(),
                            )
                        elif self.app_obj.bidding_type.GetValue() == "Proprietor":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "prop of " + self.app_obj.bidder_name.GetValue(),
                            )
                        else:
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "",
                            )
                        inline[i].text = new_text
                        count -= 1
        self.undertaking_on_going_work_info.save("undertaking_on_going_work_gen.docx")

    def undertaking_bid_validity_doc_process(self):
        count = 6
        for paragraph in self.undertaking_bid_validity_info.paragraphs:
            if count == 0:
                break
            if "<name_of_contractor>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_contractor>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_contractor>",
                            '"' + self.app_obj.contractor_name.GetValue() + '"',
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_address>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_address>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_address>",
                            self.app_obj.bidder_address.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bid_validity>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bid_validity>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bid_validity>", " ".join(self.bid_validity)
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<temp_bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<temp_bidder_name>" in inline[i].text:
                        if self.app_obj.bidding_type.GetValue() == "Partnership":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "Authorised Signatory and Partner of "
                                + self.app_obj.bidder_name.GetValue(),
                            )
                        elif self.app_obj.bidding_type.GetValue() == "Proprietor":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "prop of " + self.app_obj.bidder_name.GetValue(),
                            )
                        else:
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "",
                            )
                        inline[i].text = new_text
                        count -= 1
        self.undertaking_bid_validity_info.save("undertaking_bid_validity_gen.docx")

    def undertaking_cash_invest_doc_process(self):
        count = 6
        for paragraph in self.undertaking_min_cash_info.paragraphs:
            if count == 0:
                break
            if "<name_of_contractor>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_contractor>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_contractor>",
                            '"' + self.app_obj.contractor_name.GetValue() + '"',
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_address>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_address>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_address>",
                            self.app_obj.bidder_address.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<cash_invest>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<cash_invest>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<cash_invest>", self.app_obj.cash_invest.GetValue()
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<temp_bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<temp_bidder_name>" in inline[i].text:
                        if self.app_obj.bidding_type.GetValue() == "Partnership":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "Authorised Signatory and Partner of "
                                + self.app_obj.bidder_name.GetValue(),
                            )
                        elif self.app_obj.bidding_type.GetValue() == "Proprietor":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "prop of " + self.app_obj.bidder_name.GetValue(),
                            )
                        else:
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "",
                            )
                        inline[i].text = new_text
                        count -= 1
        self.undertaking_min_cash_info.save("undertaking_min_cash_gen.docx")

    def declaration_doc_process(self):
        count = 2
        for paragraph in self.declaration_info.paragraphs:
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
        self.declaration_info.save("declaration_no_relatives_gen.docx")

    def auth_seek_ref_doc_process(self):
        count = 10
        for paragraph in self.auth_seek_ref_info.paragraphs:
            if count == 0:
                break
            if "<name_of_contractor>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_contractor>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_contractor>",
                            '"' + self.app_obj.contractor_name.GetValue() + '"',
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_name>",
                            '"' + self.app_obj.bidder_name.GetValue() + '"',
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_address>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_address>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_address>",
                            self.app_obj.bidder_address.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<employer_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<employer_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<employer_name>",
                            self.app_obj.tender_employer_name.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_ac_num>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_ac_num>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_ac_num>",
                            self.app_obj.bidder_ac_num.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_bank_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_bank_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_bank_name>",
                            self.app_obj.bidder_bank_name.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<bidder_bank_branch>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<bidder_bank_branch>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<bidder_bank_branch>",
                            self.app_obj.bidder_bank_branch.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<temp_bidder_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<temp_bidder_name>" in inline[i].text:
                        if self.app_obj.bidding_type.GetValue() == "Partnership":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "Authorised Signatory and Partner of "
                                + self.app_obj.bidder_name.GetValue(),
                            )
                        elif self.app_obj.bidding_type.GetValue() == "Proprietor":
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "prop of " + self.app_obj.bidder_name.GetValue(),
                            )
                        else:
                            new_text = inline[i].text.replace(
                                "<temp_bidder_name>",
                                "",
                            )
                        inline[i].text = new_text
                        count -= 1
        self.auth_seek_ref_info.save("auth_seek_ref_gen.docx")

    def dispute_review_doc_process(self):
        count = 3
        for paragraph in self.dispute_review_info.paragraphs:
            if count == 0:
                break
            if "<employer_name>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<employer_name>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<employer_name>",
                            self.app_obj.tender_employer_name.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<name_of_work>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<name_of_work>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<name_of_work>",
                            self.app_obj.tender_name_of_work.GetValue(),
                        )
                        inline[i].text = new_text
                        count -= 1
            if "<date>" in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace(
                            "<date>",
                            self.date,
                        )
                        inline[i].text = new_text
                        count -= 1
        self.dispute_review_info.save("acceptance_review_dispute_expert_gen.docx")

    def milestone_doc_process(self):
        for p in self.milestone_info.paragraphs:
            if "<date>" in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if "<date>" in inline[i].text:
                        new_text = inline[i].text.replace("<date>", self.date)
                        inline[i].text = new_text
                break
        self.milestone_info.save("milestone_gen.docx")

    def combine_doc(self):
        num_of_files = len(self.list_of_prefs)
        try:
            master_file = Document(self.list_of_prefs[0])
            merged_doc_composer = Composer(master_file)
            for i in range(1, num_of_files):
                doc_temp = Document(self.list_of_prefs[i])
                merged_doc_composer.append(doc_temp)
            merged_doc_composer.save(self.path + self.merged_doc_file_name)
            try:
                for file in self.list_of_prefs:
                    if file.endswith("_gen.docx"):
                        os.remove(file)
            except OSError:
                pass
            if self.choice is False:
                self.status = True
        except IOError:
            self.status = False

    def pdf_generate(self):
        try:
            out_file = self.merged_doc_file_name.replace("docx", "pdf")
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            src_doc = word.Documents.Open(self.path + self.merged_doc_file_name)
            src_doc.SaveAs(self.path + out_file, FileFormat=17)
            src_doc.Close()
            word.Quit()
            self.status = True
        except Exception as e:
            print(e)
            self.status = False
